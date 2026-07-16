import io
import re
from pathlib import Path

import pymupdf
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def clean_text(text: str) -> str:
    """
    Clean extracted resume text while preserving line structure.
    """

    if not text:
        return ""

    # Replace Windows-style line endings.
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Replace tabs with spaces.
    text = text.replace("\t", " ")

    # Remove repeated spaces without removing line breaks.
    text = re.sub(r"[ ]{2,}", " ", text)

    # Remove spaces at the beginning and end of each line.
    lines = [line.strip() for line in text.split("\n")]

    # Remove excessive blank lines.
    cleaned_lines = []
    previous_line_blank = False

    for line in lines:
        is_blank = not line

        if is_blank and previous_line_blank:
            continue

        cleaned_lines.append(line)
        previous_line_blank = is_blank

    return "\n".join(cleaned_lines).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file supplied as bytes.
    """

    extracted_pages = []

    try:
        with pymupdf.open(stream=file_bytes, filetype="pdf") as document:
            for page in document:
                page_text = page.get_text("text", sort=True)
                extracted_pages.append(page_text)

    except Exception as error:
        raise ValueError(f"Unable to read the PDF file: {error}") from error

    return clean_text("\n".join(extracted_pages))


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract paragraphs and table content from a DOCX file.
    """

    extracted_content = []

    try:
        document = Document(io.BytesIO(file_bytes))

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                extracted_content.append(paragraph_text)

        # Resume skills and education are sometimes placed inside tables.
        for table in document.tables:
            for row in table.rows:
                row_content = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        row_content.append(cell_text)

                if row_content:
                    extracted_content.append(" | ".join(row_content))

    except Exception as error:
        raise ValueError(f"Unable to read the DOCX file: {error}") from error

    return clean_text("\n".join(extracted_content))


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """
    Choose the correct parser according to the uploaded extension.
    """

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))

        raise ValueError(
            f"Unsupported file type '{extension}'. "
            f"Supported formats: {supported}"
        )

    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    if extension == ".pdf":
        return extract_text_from_pdf(file_bytes)

    if extension == ".docx":
        return extract_text_from_docx(file_bytes)

    raise ValueError("Unsupported resume format.")